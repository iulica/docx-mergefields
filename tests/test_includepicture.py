import unittest
from os import path

import docx

from mergefields import MergeFieldsDocument


class IncludePictureFieldsTest(unittest.TestCase):
    """
    Testing local images, link images and base64 images
    """

    def test_images(self):
        """
        test a docx with 3 images without tables
        """
        # @TODO split the three kinds of INCLUDEPICTURE into separate tests
        with MergeFieldsDocument(
            path.join(path.dirname(__file__), "test_includepicture.docx")
        ) as document:
            # self.assertEqual(document.get_merge_fields(),
            #                  set(['rowno', 'url']))
            document.transform_fields()
            # document.doc.save('../docx-mailmerge/tests/output/test_includepicture.docx')
            rels = {}
            for r in document.doc.part.rels.values():
                if isinstance(r._target, docx.ImagePart):
                    rels[r.rId] = path.basename(r._target.partname)
            self.assertEqual(len(rels), 3)
            # print(rels)
        # document = docx.Document("tests/output/test_includepicture_1.docx")

    def test_table_images(self):
        """
        test a docx with 3 images in a table
        """
        # @TODO split the three kinds of INCLUDEPICTURE into separate tests
        with MergeFieldsDocument(
            path.join(path.dirname(__file__), "test_includepicture_table.docx")
        ) as document:
            # self.assertEqual(document.get_merge_fields(),
            #                  set(['rowno', 'url']))
            document.transform_fields()
            # document.doc.save('../docx-mailmerge/tests/output/test_includepicture.docx')
            rels = {}
            for r in document.doc.part.rels.values():
                if isinstance(r._target, docx.ImagePart):
                    rels[r.rId] = path.basename(r._target.partname)
            self.assertEqual(len(rels), 3)
            # print(rels)
        # document = docx.Document("tests/output/test_includepicture_1.docx")
