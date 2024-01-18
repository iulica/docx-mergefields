

import shlex
import warnings
import pathlib
from io import BytesIO
from urllib.parse import urlparse
import urllib.request
import docx
from docx.shared import Pt

NAMESPACES = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
    'ct': 'http://schemas.openxmlformats.org/package/2006/content-types',
    'rr': 'http://schemas.openxmlformats.org/package/2006/relationships',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'xml': 'http://www.w3.org/XML/1998/namespace'
}

class IncludePictureField():

    def __init__(self, field_type, instr, tokens, field_dict):
        self.field_type = field_type
        self.instr = instr
        self.tokens = tokens
        self.field_dict = field_dict
        self.image_r_elem = None
        self.flags = {}
        self._parse_args()

    def _parse_args(self):
        args = self.tokens[2:]
        last_flag_options = None
        while args:
            arg = args.pop(0)
            flag = arg[0:2]
            if not flag: # why this check ?
                continue
            if len(flag) == 2 and flag[0] == '\\':
                # we have a \ flag
                last_flag_options = self.flags.setdefault(flag[1], [])
                if arg[2:]: # no space after the flag
                    # separate the flag from the argument
                    last_flag_options.append(arg[2:])
            elif last_flag_options is not None:
                last_flag_options.append(arg)

    def _get_int_flag(self, flag, default=None):
        v_list = self.flags.get(flag)
        if v_list:
            if len(v_list) > 1:
                warnings.warn('Multiple values for flag {}: {}'.format(flag, v_list))
            try:
                return int(v_list[0])
            except ValueError:
                warnings.warn('Invalid flag value {}: {}'.format(flag, v_list))
        return default

    def _find_picture(self, doc_path):
        path = self.tokens[1]
        url_parts = urlparse(path)
        if url_parts.scheme: # in ['http', 'https', 'data']:
            image_from_url = urllib.request.urlopen(path)
            io_url = BytesIO(image_from_url.read())
            return io_url
        if doc_path is not None:
            path = doc_path / path
        return str(path)

    def insert_picture(self, doc, doc_path=None):
        first_run = self.field_dict['instr_elements'][0]
        parent = self.field_dict['parent']
        p_obj = self.find_paragraph(doc, parent)
        r_obj = self.find_run(p_obj, first_run)
        # find the paragraph and the run
        if r_obj and p_obj:
            width = self._get_int_flag('w')
            height = self._get_int_flag('h')
            if width:
                width = Pt(width)
            if height:
                height = Pt(height)

            image_path_or_stream = self._find_picture(doc_path)
            r_obj.clear()
            r_obj.add_picture(image_path_or_stream, width=width, height=height)
            self.image_r_elem = first_run

    def clean(self):
        parent = self.field_dict['parent']
        for elem in reversed(self.field_dict['all_elements']):
            if elem != self.image_r_elem:
                parent.remove(elem)

    def find_paragraph(self, doc, p_element):
        p_list = [p for p in doc.paragraphs if p._p == p_element]
        if len(p_list) == 1:
            return p_list[0]
        assert not p_list, "more than 1 paragraph found"
        warnings.warn("paragraph not found")
        return None

    def find_run(self, paragraph, r_element):
        r_list = [p for p in paragraph.runs if p._r == r_element]
        if len(r_list) == 1:
            return r_list[0]
        assert not r_list, "more than 1 runs found"
        warnings.warn("run not found")
        return None


class MailmergeDocument():

    def __init__(self, path):
        self.doc = docx.Document(path)
        self.path = pathlib.Path(path)
        self.fields = []
        self.__fill_complex_fields(self.doc._body._element)

    def __get_next_element(self, current_element):
        """ returns the next element of a complex field """
        next_element = current_element.getnext()
        current_paragraph = current_element.getparent()
        # we search through paragraphs for the next <w:r> element
        while next_element is None:
            current_paragraph = current_paragraph.getnext()
            if current_paragraph is None:
                return None, None, None
            next_element = current_paragraph.find('w:r', namespaces=NAMESPACES)

        # print(''.join(next_element.xpath('w:instrText/text()', namespaces=NAMESPACES)))
        field_char_subelem = next_element.find('w:fldChar', namespaces=NAMESPACES)
        if field_char_subelem is None:
            return next_element, None, None

        return next_element, field_char_subelem, field_char_subelem.xpath('@w:fldCharType', namespaces=NAMESPACES)[0]

    def _pull_next_merge_field(self, elements_of_type_begin, nested=False):

        assert (elements_of_type_begin)
        current_element = elements_of_type_begin.pop(0)
        parent_element = current_element.getparent()
        all_elements = [] # we need all the elments in case of updates
        instr_elements = [] # the instruction part, elements that define how to get the value
        show_elements = [] # the elements showing the current value

        current_element_list = instr_elements
        all_elements.append(current_element)
        
        # good_elements = []
        # ignore_elements = [current_element]
        # current_element_list = good_elements
        field_char_type = None
        contains_nested_fields = False
        # print('>>>>>>>')
        while field_char_type != 'end':
            # find next sibling
            next_element, field_char_subelem, field_char_type = \
                self.__get_next_element(current_element)

            if next_element is None:
                instr_text = self.get_instr_text(instr_elements, recursive=False)
                raise ValueError("begin without end near:" + instr_text)

            if field_char_type == 'begin':
                # nested elements
                contains_nested_fields = True
                assert(elements_of_type_begin[0] is next_element)
                merge_field_sub_obj, next_element = self._pull_next_merge_field(elements_of_type_begin, nested=True)
                # if merge_field_sub_obj:
                #     next_element = merge_field_sub_obj.insert_into_tree()
                # # print("current list is ignore", current_element_list is ignore_elements)
                # # print("<<<<< #####", etree.tostring(next_element))
            elif field_char_type == 'separate':
                current_element_list = show_elements
            elif next_element.tag == 'MergeField':
                assert 0, "nested simple field"
                # # we have a nested simple Field - mark it as nested

                # self.merge_data.mark_field_as_nested(next_element.get('merge_key'))

            if field_char_type not in ['end', 'separate']:
                current_element_list.append(next_element)
            all_elements.append(next_element)
            current_element = next_element

        # print('<<<<<<<', len(good_elements), len(ignore_elements))
        merge_obj = dict(
            parent=parent_element,
            nested=nested,
            contains_nested_fields=contains_nested_fields,
            all_elements=all_elements,
            instr_elements=instr_elements,
            show_elements=show_elements)
        return merge_obj, current_element

    def _get_instr_tokens(self, instr):
        s = shlex.shlex(instr, posix=True)
        s.whitespace_split = True
        s.commenters = ''
        s.escape = ''
        return s

    def _get_field_type(self, instr):
        s = shlex.split(instr, posix=False)
        return s[0], s[1:]

    def __fill_complex_fields(self, part):
        """ finds all begin fields and then builds the MergeField objects and inserts the replacement Elements in the tree """
        # will find all "runs" containing an element of fldChar type=begin
        elements_of_type_begin = list(part.xpath('.//w:r/w:fldChar[@w:fldCharType="begin"]/..'))
        # print(elements_of_type_begin)
        while elements_of_type_begin:
            merge_field_dict, _ = self._pull_next_merge_field(elements_of_type_begin)

            if merge_field_dict:
                instr = self.get_instr_text(merge_field_dict['instr_elements'])
                field_type, rest = self._get_field_type(instr)

                if field_type.upper() not in ['INCLUDEPICTURE']:
                    continue

                if merge_field_dict['nested']:
                    warnings.warn("Ignore nested fields: {}".format(instr))
                    continue

                if merge_field_dict['contains_nested_fields']:
                    warnings.warn("Ignore fields containing nested fields: {}".format(instr))
                    continue

                try:
                    tokens = list(self._get_instr_tokens(instr))
                except ValueError as e:
                    tokens = [field_type] + list(map(lambda part:part.replace('"', ''), rest))
                    warnings.warn("Invalid field description <{}> near: <{}>".format(str(e), instr))

                field_obj = IncludePictureField(field_type, instr, tokens, merge_field_dict)
                self.fields.append(field_obj)

    def get_instr_text(self, elements, recursive=False):
        return "".join([
            text
            for elem in elements
            for text in elem.xpath('w:instrText/text()') + [
                "{{{}}}".format(obj_name)
                if not recursive else
                "unknown"
                for obj_name in elem.xpath('@merge_key')]
        ])

    def transform_fields(self):
        for field in self.fields:
            # print(field.instr, field.tokens)
            field.insert_picture(self.doc, doc_path=self.path.parent)

        for field in self.fields:
            field.clean()

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        pass

if __name__ == '__main__':
    DOCX_PATH = "../docx-mailmerge/tests/test_includepicture.docx"
    with MailmergeDocument(DOCX_PATH) as mailmerge_doc:
        mailmerge_doc.transform_fields()
        mailmerge_doc.doc.save('../docx-mailmerge/tests/output/test_includepicture.docx')
